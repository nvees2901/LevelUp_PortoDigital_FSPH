"""
document.py — Service de extração de texto de documentos

Responsável por receber um arquivo (bytes) e retornar o texto extraído.
Suporta: PDF (.pdf), Word moderno (.docx), Word legado (.doc).

Por que este service existe separado?
  A lógica de extração é independente de framework.
  Pode ser testada com arquivos reais sem subir o servidor FastAPI.
  Os routes só chamam DocumentService.extract() — não conhecem pdfplumber.

Hierarquia de qualidade de extração:
  1. .docx → python-docx  (perfeito: XML estruturado)
  2. .pdf  → pdfplumber   (bom: preserva layout e tabelas)
  3. .doc  → OLE binary    (parse do Word Binary Format via olefile)
  4. .doc  → antiword      (fallback CLI, requer instalação no sistema)
"""

import io
import re
import struct
import subprocess
import tempfile

import olefile
import pdfplumber
from docx import Document as DocxDocument

from app.utils.exceptions import DocumentProcessingError, UnsupportedFormatError
from app.utils.logging import get_logger

# Magic bytes do formato OLE Compound Binary File (usado por .doc antigos)
_OLE_MAGIC = b"\xd0\xcf\x11\xe0"

logger = get_logger(__name__)


class DocumentService:

    # Extensões suportadas — usadas na validação do route
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}
    SUPPORTED_MIME_TYPES = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }

    @classmethod
    def _detect_real_format(cls, file_bytes: bytes, extension: str) -> str:
        """
        Detecta o formato real do arquivo pelos magic bytes.
        Alguns .docx são na verdade .doc (OLE) renomeados.
        """
        if file_bytes[:4] == _OLE_MAGIC and extension == ".docx":
            logger.info("Arquivo .docx é na verdade OLE (.doc) — usando extrator de .doc")
            return ".doc"
        return extension

    @classmethod
    def extract_text_sync(cls, file_bytes: bytes, filename: str) -> str:
        """
        Versão síncrona de extract_text — para uso em threads (ex: RagService).
        Os extratores internos (_extract_pdf, _extract_docx, _extract_doc) são
        todos síncronos, portanto não há necessidade de event loop.
        """
        extension = cls._get_extension(filename)
        extension = cls._detect_real_format(file_bytes, extension)

        logger.info("Extraindo texto: filename=%r ext=%s size=%d bytes",
                    filename, extension, len(file_bytes))

        if extension == ".pdf":
            return cls._extract_pdf(file_bytes, filename)
        elif extension == ".docx":
            return cls._extract_docx(file_bytes, filename)
        elif extension == ".doc":
            return cls._extract_doc(file_bytes, filename)
        else:
            raise UnsupportedFormatError(filename)

    @classmethod
    async def extract_text(cls, file_bytes: bytes, filename: str) -> str:
        """
        Extrai texto de um documento baseado na extensão do arquivo.

        Args:
            file_bytes: Conteúdo binário do arquivo (recebido via upload)
            filename:   Nome original do arquivo (usado para detectar extensão)

        Returns:
            Texto extraído como string (pode conter quebras de linha)

        Raises:
            UnsupportedFormatError: se a extensão não for suportada
            DocumentProcessingError: se a extração falhar (arquivo corrompido)
        """
        extension = cls._get_extension(filename)
        extension = cls._detect_real_format(file_bytes, extension)

        logger.info("Extraindo texto: filename=%r ext=%s size=%d bytes",
                    filename, extension, len(file_bytes))

        if extension == ".pdf":
            return cls._extract_pdf(file_bytes, filename)
        elif extension == ".docx":
            return cls._extract_docx(file_bytes, filename)
        elif extension == ".doc":
            return cls._extract_doc(file_bytes, filename)
        else:
            raise UnsupportedFormatError(filename)

    # ------------------------------------------------------------------ #
    # Extratores por formato
    # ------------------------------------------------------------------ #

    @staticmethod
    def _extract_pdf(file_bytes: bytes, filename: str) -> str:
        """
        Extrai texto de PDF usando pdfplumber.

        Por que pdfplumber?
          - Extrai tabelas como listas estruturadas (melhor que PyPDF2)
          - Preserva posicionamento do texto
          - Handles PDFs digitalizados com OCR quando combinado com pytesseract

        Tabelas são convertidas em texto formatado com separadores | para
        preservar a estrutura ao realizar análise posterior.
        """
        try:
            text_parts = []

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    # Extrai texto da página
                    page_text = page.extract_text() or ""

                    # Extrai tabelas e as converte para texto
                    tables = page.extract_tables()
                    table_texts = []
                    for table in tables:
                        for row in table:
                            # Filtra células None e une com |
                            row_text = " | ".join(
                                str(cell).strip() if cell else ""
                                for cell in row
                            )
                            if row_text.strip():
                                table_texts.append(row_text)

                    page_content = page_text
                    if table_texts:
                        page_content += "\n" + "\n".join(table_texts)

                    if page_content.strip():
                        text_parts.append(f"[Página {page_num}]\n{page_content}")

            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                raise DocumentProcessingError(
                    filename,
                    "PDF parece estar vazio ou ser uma imagem sem OCR. "
                    "Verifique se o PDF contém texto selecionável."
                )

            logger.info("PDF extraído: filename=%r páginas=%d chars=%d",
                        filename, len(text_parts), len(full_text))
            return full_text

        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(filename, str(e)) from e

    @staticmethod
    def _extract_docx(file_bytes: bytes, filename: str) -> str:
        """
        Extrai texto de arquivos .docx usando python-docx.

        O formato .docx é XML estruturado — python-docx parseia os elementos
        nativamente: paragraphs, tables, headers, footers.

        Fallback: alguns .docx têm content type não-padrão (ex: themeManager+xml)
        que o python-docx rejeita. Nesses casos, extraímos o texto diretamente
        do XML dentro do ZIP.
        """
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            text_parts = []

            # Parágrafos normais
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Tabelas — converte células para texto
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells
                        if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            full_text = "\n".join(text_parts)

            if not full_text.strip():
                raise DocumentProcessingError(filename, "Documento DOCX está vazio.")

            logger.info("DOCX extraído: filename=%r parágrafos=%d chars=%d",
                        filename, len(text_parts), len(full_text))
            return full_text

        except DocumentProcessingError:
            raise
        except Exception as e:
            logger.warning(
                "python-docx falhou para %r (%s), tentando fallback XML...",
                filename, e,
            )
            return DocumentService._extract_docx_fallback(file_bytes, filename)

    @staticmethod
    def _extract_docx_fallback(file_bytes: bytes, filename: str) -> str:
        """
        Fallback: extrai texto diretamente dos XMLs dentro do ZIP do .docx.

        Útil para arquivos com content type não-padrão que o python-docx rejeita.
        Lê word/document.xml e quaisquer outros XMLs em word/ para extrair
        todo o texto disponível.
        """
        import re
        import zipfile

        try:
            zf = zipfile.ZipFile(io.BytesIO(file_bytes))
        except zipfile.BadZipFile:
            raise DocumentProcessingError(
                filename, "Arquivo não é um DOCX válido (ZIP corrompido)."
            )

        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        text_parts: list[str] = []

        xml_targets = [n for n in zf.namelist() if n.startswith("word/") and n.endswith(".xml")]
        # Prioriza word/document.xml
        xml_targets.sort(key=lambda n: (0 if "document" in n else 1, n))

        for xml_name in xml_targets:
            try:
                from xml.etree import ElementTree as ET
                tree = ET.fromstring(zf.read(xml_name))
                for t_elem in tree.iter(f"{ns}t"):
                    if t_elem.text:
                        text_parts.append(t_elem.text)
            except ET.ParseError:
                continue

        zf.close()

        full_text = " ".join(text_parts)
        # Normaliza espaços duplicados e quebras
        full_text = re.sub(r" {2,}", " ", full_text)

        if not full_text.strip():
            raise DocumentProcessingError(filename, "Não foi possível extrair texto do DOCX.")

        logger.info(
            "DOCX extraído (fallback XML): filename=%r chars=%d",
            filename, len(full_text),
        )
        return full_text

    @staticmethod
    def _extract_doc(file_bytes: bytes, filename: str) -> str:
        """
        Extrai texto de arquivos .doc (Word legado).

        Estratégia em camadas:
          1. olefile — parse nativo do OLE, extrai texto do stream WordDocument
          2. antiword — CLI (requer instalação, disponível no Docker Linux)
        """
        # Tenta olefile primeiro (funciona em qualquer OS, sem dependência externa)
        try:
            text = DocumentService._extract_doc_olefile(file_bytes, filename)
            if text and text.strip():
                return text
        except Exception as e:
            logger.debug("olefile falhou para %r: %s", filename, e)

        # Fallback: antiword (Linux/Docker)
        try:
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name

            result = subprocess.run(
                ["antiword", tmp_path],
                capture_output=True,
                text=True,
                timeout=30,
            )

            if result.returncode == 0 and result.stdout.strip():
                logger.info("DOC extraído via antiword: filename=%r chars=%d",
                            filename, len(result.stdout))
                return result.stdout

        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        except Exception as e:
            logger.debug("antiword falhou para %r: %s", filename, e)

        raise DocumentProcessingError(
            filename,
            "Não foi possível extrair texto do arquivo .doc. "
            "Converta para .docx e tente novamente."
        )

    @staticmethod
    def _extract_doc_olefile(file_bytes: bytes, filename: str) -> str:
        """
        Extrai texto de .doc via olefile — parse do Word Binary Format.

        O texto no formato .doc fica no stream 'WordDocument' e '1Table'/'0Table'.
        Para simplificar, extraímos bytes de texto Unicode do stream WordDocument.
        """
        ole = olefile.OleFileIO(io.BytesIO(file_bytes))

        text_parts = []

        # Tenta extrair do stream WordDocument (contém o texto principal)
        if ole.exists("WordDocument"):
            word_stream = ole.openstream("WordDocument").read()

            # Extrai strings UTF-16LE do stream binário
            # O texto do Word Binary fica intercalado com metadados
            decoded = DocumentService._extract_strings_from_binary(word_stream)
            if decoded:
                text_parts.append(decoded)

        # Também verifica se há ObjectPool com textos adicionais
        for stream_name in ole.listdir():
            name = "/".join(stream_name)
            if any(x in name.lower() for x in ["compobj", "summary", "table"]):
                continue
            if name == "WordDocument":
                continue
            try:
                data = ole.openstream(stream_name).read()
                extra = DocumentService._extract_strings_from_binary(data)
                if extra and len(extra) > 50:
                    text_parts.append(extra)
            except Exception:
                continue

        ole.close()

        full_text = "\n".join(text_parts)
        full_text = re.sub(r"\n{3,}", "\n\n", full_text)

        if full_text.strip():
            logger.info("DOC extraído via olefile: filename=%r chars=%d",
                        filename, len(full_text))

        return full_text

    @staticmethod
    def _extract_strings_from_binary(data: bytes, min_length: int = 10) -> str:
        """
        Extrai strings legíveis (UTF-16LE e ASCII) de dados binários.
        Filtra strings muito curtas para evitar ruído de metadados.
        """
        # Tenta UTF-16LE primeiro (padrão do Word)
        try:
            # Procura sequências de chars UTF-16LE printáveis
            text_parts = []
            i = 0
            current = []
            while i < len(data) - 1:
                # UTF-16LE: char seguido de 0x00 para ASCII range
                char_val = struct.unpack_from("<H", data, i)[0]
                if 0x20 <= char_val < 0xFFFE:
                    current.append(chr(char_val))
                else:
                    if len(current) >= min_length:
                        text_parts.append("".join(current))
                    current = []
                    if char_val in (0x000D, 0x000A):
                        text_parts.append("\n")
                i += 2

            if len(current) >= min_length:
                text_parts.append("".join(current))

            result = "".join(text_parts)
            # Limpa caracteres de controle residuais
            result = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", result)
            return result.strip()
        except Exception:
            return ""

    # ------------------------------------------------------------------ #
    # Utilitários
    # ------------------------------------------------------------------ #

    @staticmethod
    def _get_extension(filename: str) -> str:
        """Extrai a extensão em minúsculas do nome do arquivo."""
        parts = filename.rsplit(".", 1)
        if len(parts) < 2:
            return ""
        return f".{parts[1].lower()}"

    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """Verifica se a extensão do arquivo é suportada."""
        return cls._get_extension(filename) in cls.SUPPORTED_EXTENSIONS
