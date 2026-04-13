import chromadb
import time
import re
import ollama
from sentence_transformers import SentenceTransformer
from docx import Document
import pdfplumber
import os
import json
from dotenv import load_dotenv

load_dotenv(dotenv_path=".env")

MODO_IA = os.getenv("MODO_IA", "ollama").strip().lower()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "").strip()
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.0-flash").strip()
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen2.5:3b").strip()

if MODO_IA == "gemini" and not GEMINI_API_KEY:
    raise Exception("GEMINI_API_KEY não carregada. Verifique o .env")

client_gemini = None

if MODO_IA == "gemini" and GEMINI_API_KEY:
    try:
        from google import genai
        client_gemini = genai.Client(api_key=GEMINI_API_KEY)
        print(f"Modo GEMINI ativado ({GEMINI_MODEL})")
    except Exception as e:
        print(f"Erro ao iniciar Gemini: {e}. Usando Ollama como fallback.")
        MODO_IA = "ollama"

if MODO_IA == "ollama":
    print(f"Modo OLLAMA ativado ({OLLAMA_MODEL})")

client = chromadb.PersistentClient(path="./meu_banco_dados")
NOME_COLECAO = "projeto_hemocentro"
modelo = SentenceTransformer('all-MiniLM-L6-v2')

PASTA_DOCUMENTOS = "./documentos"

GATILHOS_RELATORIO = [
    "relatório", "relatorio", "gerar termo", "elaborar",
    "criar termo", "documento completo", "todas as seções",
    "todas as secoes", "resumo completo", "gere um", "crie um",
    "elabore um", "monte um", "produza um"
]

OPCOES_RAPIDO = {"num_predict": 200, "temperature": 0.1, "num_ctx": 1024}
OPCOES_RELATORIO = {"num_predict": 1000, "temperature": 0.2, "num_ctx": 2048}


def detectar_modo(pergunta):
    pergunta_lower = pergunta.lower()
    for gatilho in GATILHOS_RELATORIO:
        if gatilho in pergunta_lower:
            return "relatorio"
    return "rapido"


def gerar_resposta(prompt, modo):
    prompt = prompt[:12000]

    if MODO_IA == "gemini" and client_gemini:
        for tentativa in range(3):
            try:
                resposta = client_gemini.models.generate_content(
                    model=GEMINI_MODEL,
                    contents=prompt
                )
                return resposta.text
            except Exception as e:
                erro = str(e)
                if "429" in erro or "RESOURCE_EXHAUSTED" in erro:
                    espera = 5 * (tentativa + 1)
                    print(f"Limite de requisições atingido. Aguardando {espera}s (tentativa {tentativa + 1}/3)...")
                    time.sleep(espera)
                    continue
                elif "404" in erro or "NOT_FOUND" in erro:
                    return "Modelo não encontrado. Verifique GEMINI_MODEL no .env"
                else:
                    print(f"Erro Gemini (tentativa {tentativa + 1}/3): {erro}")
                    return f"Erro inesperado: {erro}"
        return "Serviço Gemini indisponível após 3 tentativas. Tente novamente mais tarde."

    try:
        opcoes = OPCOES_RELATORIO if modo == "relatorio" else OPCOES_RAPIDO
        resposta = ollama.chat(
            model=OLLAMA_MODEL,
            messages=[{"role": "user", "content": prompt}],
            options=opcoes
        )
        return resposta['message']['content']
    except Exception as e:
        return f"Erro ao conectar com Ollama: {str(e)}"


def carregar_criterios():
    try:
        with open("criterios.json", "r", encoding="utf-8") as f:
            criterios = json.load(f)
        print(f"{len(criterios)} critérios da Lei 14.133 carregados.")
        return criterios
    except FileNotFoundError:
        print("Arquivo criterios.json não encontrado.")
        return []


CRITERIOS_LEI = carregar_criterios()


def validar_criterios(texto):
    texto_lower = texto.lower()
    resultado = []
    aprovados = 0

    for criterio in CRITERIOS_LEI:
        encontrado = any(
            re.search(rf'\b{re.escape(kw.lower())}\b', texto_lower)
            for kw in criterio["keywords"]
        )
        status = "aprovado" if encontrado else "reprovado"
        if encontrado:
            aprovados += 1
        resultado.append({
            "id": criterio["id"],
            "nome": criterio["nome"],
            "artigo": criterio["artigo"],
            "descricao": criterio["descricao"],
            "status": status
        })

    score = int((aprovados / len(CRITERIOS_LEI)) * 100) if CRITERIOS_LEI else 0
    return resultado, score


def criar_chunks(texto, tamanho=100, sobreposicao=20):
    palavras = texto.split()
    chunks = []
    i = 0
    while i < len(palavras):
        chunk = " ".join(palavras[i:i + tamanho])
        if len(chunk) > 80:
            chunks.append(chunk)
        i += tamanho - sobreposicao
    return chunks


def _indexar_chunks(texto, caminho, colecao, show_progress=False):
    chunks = criar_chunks(texto)
    if not chunks:
        print(f"{caminho}: nenhum conteúdo encontrado.")
        return
    vetores = modelo.encode(chunks, batch_size=32, show_progress_bar=show_progress).tolist()
    ids = [f"{caminho}_chunk_{i}" for i in range(len(chunks))]
    metadatas = [{"fonte": caminho}] * len(chunks)
    try:
        colecao.add(embeddings=vetores, documents=chunks, ids=ids, metadatas=metadatas)
        print(f"{caminho}: {len(chunks)} chunks indexados.")
    except Exception as e:
        print(f"Erro ao salvar {caminho}: {e}")


def ler_word(caminho, colecao):
    try:
        doc = Document(caminho)
        texto_completo = "\n".join(
            p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 30
        )
        _indexar_chunks(texto_completo, caminho, colecao)
    except Exception as e:
        print(f"Erro no Word {caminho}: {e}")


def ler_pdf(caminho, colecao):
    try:
        with pdfplumber.open(caminho) as pdf:
            texto_completo = "\n".join(
                p.extract_text() for p in pdf.pages if p.extract_text()
            )
        if not texto_completo.strip():
            print(f"{caminho}: PDF sem texto extraível (pode ser escaneado).")
            return
        _indexar_chunks(texto_completo, caminho, colecao, show_progress=True)
    except Exception as e:
        print(f"Erro no PDF {caminho}: {e}")


def ler_txt(caminho, colecao):
    try:
        with open(caminho, "r", encoding="utf-8") as f:
            texto_completo = f.read()
        if not texto_completo.strip():
            print(f"{caminho}: arquivo vazio.")
            return
        _indexar_chunks(texto_completo, caminho, colecao)
    except Exception as e:
        print(f"Erro no TXT {caminho}: {e}")


def buscar(pergunta, colecao):
    modo = detectar_modo(pergunta)

    if modo == "relatorio":
        print("\nModo RELATÓRIO detectado. Gerando resposta completa, aguarde...")
    else:
        print("\nModo rápido...")

    vetor_pergunta = modelo.encode(pergunta).tolist()

    resultado = colecao.query(
        query_embeddings=[vetor_pergunta],
        n_results=5 if modo == "relatorio" else 3,
        include=["documents", "metadatas"]
    )

    if not resultado['documents'][0]:
        return json.dumps({
            "resposta": "Não encontrei essa informação nos documentos.",
            "fonte": None,
            "score": 0,
            "criterios": [],
            "modo": modo
        }, ensure_ascii=False, indent=2)

    partes = []
    fontes = []
    tamanho_acumulado = 0
    for doc, meta in zip(resultado['documents'][0], resultado['metadatas'][0]):
        nome_arquivo = os.path.basename(meta['fonte'])
        bloco = f"[{nome_arquivo}]\n{doc}"
        if tamanho_acumulado + len(bloco) > 8000:
            break
        partes.append(bloco)
        tamanho_acumulado += len(bloco)
        if nome_arquivo not in fontes:
            fontes.append(nome_arquivo)
    contexto = "\n\n".join(partes)

    criterios, score = validar_criterios(contexto)

    if modo == "relatorio":
        prompt = f"""Você é um assistente especializado em licitações e Termos de Referência da FSPH (Fundação de Saúde Parreiras Horta).
Use APENAS os trechos abaixo para gerar um relatório completo e detalhado em português brasileiro.
Organize a resposta com seções claras. Cite os artigos da Lei 14.133 quando relevante.
Se alguma informação não estiver nos trechos, indique claramente.

TRECHOS DOS DOCUMENTOS:
{contexto}

SOLICITAÇÃO: {pergunta}

RELATÓRIO:"""
    else:
        prompt = f"""Você é um assistente especializado em licitações e Termos de Referência da FSPH (Fundação de Saúde Parreiras Horta).
Use APENAS os trechos abaixo para responder.
Responda de forma única, clara e direta em português brasileiro.
Se a resposta não estiver nos trechos, diga apenas: "Não encontrei essa informação nos documentos."
Proibido inventar ou usar conhecimento externo.

TRECHOS DOS DOCUMENTOS:
{contexto}

PERGUNTA: {pergunta}

RESPOSTA:"""

    print(f"Gerando resposta via {MODO_IA.upper()}...\n")
    resposta_texto = gerar_resposta(prompt, modo)

    saida = {
        "resposta": resposta_texto,
        "fonte": fontes,
        "score": score,
        "criterios": criterios,
        "modo": modo
    }

    return json.dumps(saida, ensure_ascii=False, indent=2)


def analisar_documento(caminho_arquivo):
    print(f"\nAnalisando: {caminho_arquivo}")
    texto = ""

    if caminho_arquivo.endswith(".docx"):
        doc = Document(caminho_arquivo)
        texto = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    elif caminho_arquivo.endswith(".pdf"):
        with pdfplumber.open(caminho_arquivo) as pdf:
            texto = "\n".join(p.extract_text() for p in pdf.pages if p.extract_text())
    elif caminho_arquivo.endswith(".txt"):
        with open(caminho_arquivo, "r", encoding="utf-8") as f:
            texto = f.read()

    if not texto:
        return json.dumps({"erro": "Não foi possível extrair texto do documento."}, ensure_ascii=False)

    criterios, score = validar_criterios(texto)

    saida = {
        "arquivo": os.path.basename(caminho_arquivo),
        "score": score,
        "criterios": criterios,
        "aprovados": sum(1 for c in criterios if c["status"] == "aprovado"),
        "reprovados": sum(1 for c in criterios if c["status"] == "reprovado"),
    }

    return json.dumps(saida, ensure_ascii=False, indent=2)


def preparar_sistema():
    print("\n--- GERENCIADOR DE MEMÓRIA VETORIAL ---")
    print("[1] Atualizar banco (ler arquivos do zero)")
    print("[2] Atualização incremental (só arquivos novos)")
    print("[3] Modo rápido (usar banco existente)")
    opcao = input("Escolha uma opção: ")

    if opcao == '1':
        try:
            client.delete_collection(name=NOME_COLECAO)
            print("Memória antiga limpa.")
        except Exception:
            pass
        colecao = client.get_or_create_collection(name=NOME_COLECAO)

        if not os.path.exists(PASTA_DOCUMENTOS):
            os.makedirs(PASTA_DOCUMENTOS)
            print(f"Pasta '{PASTA_DOCUMENTOS}' criada. Adicione seus arquivos e rode novamente.")
            return colecao

        print("Lendo e indexando todos os arquivos...")
        for arquivo in os.listdir(PASTA_DOCUMENTOS):
            caminho = os.path.join(PASTA_DOCUMENTOS, arquivo)
            if arquivo.endswith(".docx"):
                ler_word(caminho, colecao)
            elif arquivo.endswith(".pdf"):
                ler_pdf(caminho, colecao)
            elif arquivo.endswith(".txt"):
                ler_txt(caminho, colecao)
        print("Banco atualizado com sucesso.")

    elif opcao == '2':
        colecao = client.get_or_create_collection(name=NOME_COLECAO)

        if not os.path.exists(PASTA_DOCUMENTOS):
            print("Pasta documentos não encontrada.")
            return colecao

        ids_existentes = colecao.get()['ids']
        novos = 0

        print("Verificando arquivos novos...")
        for arquivo in os.listdir(PASTA_DOCUMENTOS):
            caminho = os.path.join(PASTA_DOCUMENTOS, arquivo)
            ja_indexado = any(id.startswith(caminho) for id in ids_existentes)

            if not ja_indexado:
                print(f"Arquivo novo encontrado: {arquivo}")
                if arquivo.endswith(".docx"):
                    ler_word(caminho, colecao)
                elif arquivo.endswith(".pdf"):
                    ler_pdf(caminho, colecao)
                elif arquivo.endswith(".txt"):
                    ler_txt(caminho, colecao)
                novos += 1

        if novos == 0:
            print("Nenhum arquivo novo. Banco já está atualizado.")
        else:
            print(f"{novos} arquivo(s) novo(s) indexado(s).")

    else:
        colecao = client.get_or_create_collection(name=NOME_COLECAO)
        print("Banco carregado.")

    return colecao


if __name__ == "__main__":
    minha_colecao = preparar_sistema()

    print("\n" + "=" * 50)
    print("   IA HEMOCENTRO — PRONTA PARA CONSULTAS")
    print("=" * 50)

    while True:
        duvida = input("\nSua pergunta: ").strip()

        if duvida.lower() in ['sair', 'parar', 'exit']:
            print("Chat encerrado.")
            break

        if not duvida:
            continue

        if duvida.lower().startswith("analisar "):
            arquivo = duvida[9:].strip()
            caminho = os.path.join(PASTA_DOCUMENTOS, arquivo)
            resultado = analisar_documento(caminho)
            print("\n" + "-" * 50)
            print(resultado)
            print("-" * 50)
            continue

        resultado = buscar(duvida, minha_colecao)
        dados = json.loads(resultado)
        print("\n" + "-" * 50)
        print(f"Fonte: {', '.join(dados['fonte']) if dados['fonte'] else 'N/A'}")
        print(f"Score Lei 14.133: {dados['score']}%")
        print(f"Modo: {'Relatório' if dados['modo'] == 'relatorio' else 'Rápido'}")

        reprovados = [c['nome'] for c in dados['criterios'] if c['status'] == 'reprovado']
        if reprovados:
            print(f"Critérios não atendidos: {', '.join(reprovados)}")

        print(f"\nResposta:\n{dados['resposta']}")
        print("-" * 50)